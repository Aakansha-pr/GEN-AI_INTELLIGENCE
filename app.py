#              STUDYGEN AI - INTELLIGENT LEARNING ASSISTANT
# ============================================================

# ==============STEP 1: IMPORT LIBRARIES ================

import os
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
import streamlit as st
import numpy
import time
from PIL import Image
from dotenv import load_dotenv
from docx import Document
import io

#====================STEP 2 API KEYS======================
st.set_page_config(page_title = "GEN-AI INTELLIGENT LEARNING🎓",
              layout = "wide")

st.sidebar.title("SET API CONFIG")
st.title("GEN-AI INTELLIGENT LEARNING💻")

GOOGLE_API_KEY = st.sidebar.text_input("GOOGLE_API_KEY",type = "password")
os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY

if GOOGLE_API_KEY:
  st.sidebar.success("API key Loaded!!")
else:
  st.sidebar.info("Give API key")

if GOOGLE_API_KEY:
    model = ChatGoogleGenerativeAI(
        model="gemini-3.5-flash-lite",
        google_api_key=GOOGLE_API_KEY
    )
else:
    st.stop()

# =================== STEP 3 : PDF BACKEND FUNCTIONS ===================
def load_pdf(pdf_path):
    loader = PyPDFLoader(pdf_path)
    documents = loader.load()
    return documents
  
# =================== STEP 4 : CREATE TEXT CHUNKS ===================
def split_pdf(documents):
    """
    This function splits PDF into smaller text chunks.
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )

    chunks = text_splitter.split_documents(documents)

    return chunks

# ================= STEP 5 : CREATE EMBEDDINGS =================

def create_vectorstore(chunks, embeddings):
    vectorstore = FAISS.from_documents(
        chunks,
        embeddings
    )
    return vectorstore

def create_retriever(vectorstore, k_value):
    retriever = vectorstore.as_retriever(
        search_kwargs={"k": k_value}
    )
    return retriever
  
@st.cache_resource
def create_embeddings():
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )
    return embeddings

# ================= STEP 9 : CREATE LCEL RAG CHAIN =================

def create_rag_chain(retriever):
    """
    This function creates the LCEL RAG Chain.
    """
    prompt = ChatPromptTemplate.from_template(
        """
    You are StudyGen AI, an Intelligent Learning Assistant.

    Answer the user's question ONLY using the provided context.

    If the answer is not available in the context,
    reply:
    "I couldn't find this information in the uploaded study material."
    Context:
    {context}
    Question:
    {question}
    """
    )

    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)

    rag_chain = (
        {
            "context": retriever | format_docs,
            "question": RunnablePassthrough()
        }
        | prompt
        | model
        | StrOutputParser()
    )
    return rag_chain

# ================= STEP 10 : NOTES GENERATOR =================

def generate_notes(rag_chain):

    prompt = """
Generate well-structured study notes from the uploaded study material.

Include the following sections:

1. Introduction
2. Key Concepts
3. Important Definitions
4. Key Points
5. Examples (if available)
6. Summary

Use simple language.
Format the output using headings and bullet points.
"""

    try:
        notes = rag_chain.invoke(prompt)
        return notes

    except Exception as e:
        st.exception(e)
        return None
# ================= STEP 11 : QUIZ GENERATOR =================

def generate_quiz(rag_chain):
    """
    This function generates a quiz
    from the uploaded PDF.
    """

    prompt = """
Generate exactly 10 Multiple Choice Questions (MCQs) from the uploaded study material.

Follow this format STRICTLY.

Question 1:
<question>

A. Option 1
B. Option 2
C. Option 3
D. Option 4

Correct Answer:
A

Explanation:
Write the explanation in 2-3 complete sentences.

Repeat the same format for all 10 questions.

Rules:
- Put every option on a separate line.
- Leave one blank line between sections.
- Put "Correct Answer" on a separate line.
- Put "Explanation" on a separate line.
- Make the output clean, readable and well-formatted using Markdown.
"""

    quiz = rag_chain.invoke(prompt)

    return quiz
  
# ================= STEP 12 : DOUBT SOLVER =================

def solve_doubt(rag_chain, question):
    """
    This function answers the user's question
    from the uploaded study material.
    with proper spacing
    """

    answer = rag_chain.invoke(question)

    return answer


# ================= STEP 13 : STUDY PLANNER =================

# ================= STEP 13 : STUDY PLANNER =================

def generate_study_plan(subjects, exam_date, study_hours):

    prompt = f"""
You are StudyGen AI.

Create a personalized study plan.

Subjects:
{subjects}

Exam Date:
{exam_date}

Study Hours Per Day:
{study_hours}

Generate the output in proper Markdown.

Rules:
- Use headings.
- Use bullet points.
- Use tables where possible.
- Leave blank lines between sections.
- Do NOT return JSON.
- Do NOT return Python objects.
- Return only the study plan.
"""

    response = model.invoke(prompt)

    if isinstance(response.content, list):
        plan = "".join(
            item["text"] if isinstance(item, dict) else item.text
            for item in response.content
        )
    else:
        plan = response.content

    return plan
  
def create_doc(text):

    doc = Document()
    doc.add_heading("StudyGen AI", level=1)
    doc.add_paragraph(text)

    file = io.BytesIO()
    doc.save(file)
    file.seek(0)

    return file.getvalue()
# ============================================================
# ============== STEP 14 : STREAMLIT USER INTERFACE ===========
# ---------------------- PAGE TITLE -----------------------
st.title("📚 StudyGen AI")
st.subheader("Intelligent Learning Assistant")

st.markdown(
  """
Welcome to **StudyGen AI**.
Upload your study material and use AI to simplify learning.
"""
)

st.image("pg.png", use_container_width=True)

st.divider()
st.markdown("## ✨ Features")

col1, col2, col3, col4 = st.columns(4)

with col1:
    st.info("📄\n\nGenerate Notes")

with col2:
    st.info("📝\n\nGenerate Quiz")

with col3:
    st.info("💬\n\nSolve Doubts")

with col4:
    st.info("📅\n\nStudy Planner")

st.markdown("## 🎯 About StudyGen AI")

st.write(
  """
StudyGen AI is an AI-powered intelligent learning assistant developed using
Generative AI, Retrieval-Augmented Generation (RAG), LangChain, FAISS,
HuggingFace Embeddings, and Google's Gemini model.

The application helps students understand study materials more effectively by
providing AI-powered learning assistance from uploaded PDF documents.

### Key Features

✔ Generate well-structured study notes

✔ Create AI-powered multiple-choice quizzes

✔ Ask questions directly from the uploaded PDF

✔ Generate personalized study plans based on exam dates and study hours

StudyGen AI reduces manual study effort, improves learning efficiency,
and provides an interactive, personalized learning experience.
"""
)

st.divider()
# ==================== STEP 15 : SIDEBAR ====================

st.sidebar.title("📚 StudyGen AI")
st.sidebar.markdown("### Upload your study material")

uploaded_file = st.sidebar.file_uploader(
    "Choose a PDF File",
    type=["pdf"]
)
k_value = st.sidebar.slider(
    "Retriever Top-K",
    min_value=1,
    max_value=10,
    value=3
)

st.sidebar.divider()

st.sidebar.markdown("## ℹ️ About")

st.sidebar.info("""
📚 **StudyGen AI**

An AI-powered Intelligent Learning Assistant that transforms PDF study material into:

• 📄 Smart Notes

• 📝 AI Quizzes

• 💬 Doubt Resolution

• 📅 Personalized Study Plans

**Tech Stack**

🤖 Gemini

🔗 LangChain

🧠 FAISS

🤗 HuggingFace

🐍 Python

🎨 Streamlit
""")

# ================= STEP 16 : PROCESS PDF =================
def build_rag(pdf_path, k_value):

    documents = load_pdf(pdf_path)
    chunks = split_pdf(documents)

    embeddings = create_embeddings()
    vectorstore = create_vectorstore(chunks, embeddings)
    retriever = create_retriever(vectorstore, k_value)
    return create_rag_chain(retriever)

rag_chain = None
if uploaded_file is not None:

    save_dir = "uploaded_files"
    os.makedirs(save_dir, exist_ok=True)

    pdf_path = os.path.join(save_dir, uploaded_file.name)
    with open(pdf_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    try:
        with st.spinner("Processing PDF..."):
            rag_chain = build_rag(pdf_path, k_value)

        st.success("PDF Processed Successfully ✅")
        st.write("RAG Chain Built Successfully")

    except Exception as e:
        st.exception(e)
        
# ================= STEP 17 : SELECT FEATURE ==================
st.divider()

if rag_chain:

    col1, col2, col3 = st.columns(3)

    col1.metric("AI Model", "Gemini")
    col2.metric("Retriever", f"Top {k_value}")
    col3.metric("PDF", "Loaded")

    tab1, tab2, tab3, tab4 = st.tabs(
        [
            "📄 Smart Notes",
            "📝 AI Quiz Generator",
            "💬 AI Doubt Solver",
            "📅 Personalized Study Planner"
        ]
    )

    # ================= NOTES =================

    with tab1:

        st.subheader("📄 Generate Study Notes")

        if st.button("Generate Notes"):

            with st.spinner("Generating Notes..."):

                prompt = """
Generate well-structured study notes from the uploaded study material.

Include the following sections:

1. Introduction
2. Key Concepts
3. Important Definitions
4. Key Points
5. Examples (if available)
6. Summary

Use simple language.
Format the output using headings and bullet points.
"""

                notes = rag_chain.invoke(prompt)

                st.markdown(notes)

                st.download_button(
                    "⬇ Download Notes (.md)",
                    notes,
                    file_name="Study_Notes.md",
                    mime="text/markdown"
                )

                st.download_button(
                    "⬇ Download Notes (.docx)",
                    create_doc(notes),
                    file_name="Study_Notes.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # ================= QUIZ =================

    with tab2:

        st.subheader("📝 Generate Quiz")

        if st.button("Generate Quiz"):

            with st.spinner("Generating Quiz..."):

                prompt = """
Generate exactly 10 Multiple Choice Questions (MCQs) from the uploaded study material.

Follow this format STRICTLY.

Question 1:

<question>

A. Option 1
B. Option 2
C. Option 3
D. Option 4

Correct Answer:
A

Explanation:
Write the explanation in 2-3 sentences.

Repeat for all 10 questions.
"""

                quiz = rag_chain.invoke(prompt)

                st.markdown(quiz)

                st.download_button(
                    "⬇ Download Quiz (.md)",
                    quiz,
                    file_name="Quiz.md",
                    mime="text/markdown"
                )

                st.download_button(
                    "⬇ Download Quiz (.docx)",
                    create_doc(quiz),
                    file_name="Quiz.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # ================= DOUBT SOLVER =================

    with tab3:

        st.subheader("💬 AI Doubt Solver")

        question = st.text_input("Ask your question")

        if st.button("Get Answer"):

            with st.spinner("Generating Answer..."):

                answer = rag_chain.invoke(question)

                st.markdown(answer)

                st.download_button(
                    "⬇ Download Answer (.md)",
                    answer,
                    file_name="Answer.md",
                    mime="text/markdown"
                )

                st.download_button(
                    "⬇ Download Answer (.docx)",
                    create_doc(answer),
                    file_name="Answer.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # ================= STUDY PLANNER =================

    with tab4:
    
        st.subheader("📅 Personalized Study Planner")
    
        subjects = st.text_input("Subjects")
    
        exam_date = st.date_input("Exam Date")
    
        study_hours = st.slider(
            "Study Hours",
            1,
            12,
            4
        )
    
        if st.button("Generate Study Plan"):
    
            plan = generate_study_plan(
                subjects,
                exam_date,
                study_hours
            )
    
            # Display the formatted study plan
            st.write(plan)
    
            # Download as Markdown
            st.download_button(
                label="⬇ Download Study Plan (.md)",
                data=plan,
                file_name="Study_Plan.md",
                mime="text/markdown"
            )
    
            # Download as Word
            st.download_button(
                label="⬇ Download Study Plan (.docx)",
                data=create_doc(plan),
                file_name="Study_Plan.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
st.divider()

st.caption(
    "🚀 Powered by Gemini 2.5 Flash Lite | LangChain | FAISS | HuggingFace Embeddings | Streamlit"
)
